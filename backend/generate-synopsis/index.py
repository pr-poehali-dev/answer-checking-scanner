"""
Генерация конспекта урока через GigaChat → возвращает DOCX файл.
POST / body: {subject, class_num, topic, description, teacher_name, teacher_school, login?}
Возвращает: {docx_b64, filename, word_count, topic, subject, class_num}
"""
import json
import os
import io
import re
import ssl
import uuid
import base64
import urllib.parse
import urllib.request
import urllib.error
from datetime import datetime, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor as DocxRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTH_URL = os.environ.get("AUTH_FUNCTION_URL", "https://functions.poehali.dev/b08ae7cf-6c0b-4178-acc9-4b62b2c2a61b")
TOKENS_COST_SYNOPSIS = 5000

CORS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
}


def _resp(status: int, body: dict) -> dict:
    return {
        "statusCode": status,
        "headers": {**CORS, "Content-Type": "application/json"},
        "body": json.dumps(body, ensure_ascii=False),
        "isBase64Encoded": False,
    }


# ─── ТОКЕНЫ ПОЛЬЗОВАТЕЛЯ ──────────────────────────────────────────────────────

def spend_ai_tokens(login: str, amount: int) -> tuple[bool, str]:
    if not login:
        return True, ""
    try:
        req = urllib.request.Request(
            f"{AUTH_URL}?action=spend-tokens",
            data=json.dumps({"login": login, "amount": amount}).encode("utf-8"),
            method="POST",
            headers={"Content-Type": "application/json"},
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            json.loads(r.read().decode())
        return True, ""
    except urllib.error.HTTPError as e:
        err_body = {}
        try:
            err_body = json.loads(e.read().decode())
        except Exception:
            pass
        if e.code == 402:
            return False, err_body.get("error", "Недостаточно токенов")
        return True, ""
    except Exception:
        return True, ""


# ─── GIGACHAT AUTH ────────────────────────────────────────────────────────────

_TOKEN_CACHE: dict = {"token": None, "expires_at": None}


def _ssl_ctx():
    ctx = ssl.create_default_context()
    ctx.check_hostname = False
    ctx.verify_mode = ssl.CERT_NONE
    return ctx


def get_gigachat_token() -> str:
    now = datetime.utcnow()
    if _TOKEN_CACHE["token"] and _TOKEN_CACHE["expires_at"] and _TOKEN_CACHE["expires_at"] > now:
        return _TOKEN_CACHE["token"]
    auth_key = os.environ.get("GIGACHAT_AUTH_KEY", "").strip()
    if not auth_key:
        raise RuntimeError("GIGACHAT_AUTH_KEY не задан")
    rq_uid = str(uuid.uuid4())
    data = urllib.parse.urlencode({"scope": "GIGACHAT_API_PERS"}).encode()
    req = urllib.request.Request(
        "https://ngw.devices.sberbank.ru:9443/api/v2/oauth",
        data=data, method="POST",
        headers={
            "Content-Type": "application/x-www-form-urlencoded",
            "Accept": "application/json",
            "RqUID": rq_uid,
            "Authorization": f"Basic {auth_key}",
        },
    )
    with urllib.request.urlopen(req, timeout=20, context=_ssl_ctx()) as r:
        body = json.loads(r.read().decode())
    token = body.get("access_token")
    if not token:
        raise RuntimeError(f"GigaChat не вернул access_token: {body}")
    expires_in_ms = body.get("expires_at")
    expires_at = (datetime.utcfromtimestamp(expires_in_ms / 1000) - timedelta(minutes=2)
                  if expires_in_ms else now + timedelta(minutes=25))
    _TOKEN_CACHE["token"] = token
    _TOKEN_CACHE["expires_at"] = expires_at
    return token


def gigachat_chat(messages: list, max_tokens: int = 8000, req_timeout: int = 120) -> str:
    last_err = None
    for model in ("GigaChat-2-Max", "GigaChat-2", "GigaChat"):
        try:
            token = get_gigachat_token()
            payload = {
                "model": model,
                "messages": messages,
                "temperature": 0.4,
                "max_tokens": max_tokens,
                "stream": False,
            }
            req = urllib.request.Request(
                "https://gigachat.devices.sberbank.ru/api/v1/chat/completions",
                data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
                method="POST",
                headers={
                    "Content-Type": "application/json",
                    "Accept": "application/json",
                    "Authorization": f"Bearer {token}",
                    "Connection": "close",
                },
            )
            with urllib.request.urlopen(req, timeout=req_timeout, context=_ssl_ctx()) as r:
                body = json.loads(r.read().decode())
            choices = body.get("choices") or []
            if not choices:
                last_err = RuntimeError(f"GigaChat пустой ответ ({model}): {body}")
                continue
            content = choices[0].get("message", {}).get("content", "").strip()
            if not content:
                last_err = RuntimeError(f"GigaChat пустой content ({model})")
                continue
            return content
        except urllib.error.HTTPError as e:
            err_text = e.read().decode(errors="ignore")[:300]
            if e.code in (401, 403):
                _TOKEN_CACHE["token"] = None
                raise RuntimeError(f"GigaChat auth error {e.code}: {err_text}")
            if e.code == 404:
                last_err = RuntimeError(f"GigaChat модель недоступна ({model}): {err_text}")
                continue
            last_err = RuntimeError(f"GigaChat HTTP {e.code}: {err_text}")
        except Exception as e:
            last_err = RuntimeError(f"GigaChat недоступен ({model}): {e}")
            _TOKEN_CACHE["token"] = None
    raise last_err or RuntimeError("GigaChat: не удалось получить ответ")


# ─── ГЕНЕРАЦИЯ ТЕКСТА КОНСПЕКТА ───────────────────────────────────────────────

def generate_synopsis_text(subject: str, class_num: int, topic: str, description: str,
                            teacher_name: str, teacher_school: str) -> str:
    system = (
        "Ты — опытный учитель-методист с 20-летним стажем, эксперт по ФГОС и программам "
        "Министерства просвещения РФ. Создавай профессиональные, ДЕТАЛЬНЫЕ конспекты уроков. "
        "Раздел «Изучение нового материала» должен содержать ПОЛНОЕ объяснение темы "
        "с примерами, формулами, датами, историческими фактами. "
        "Минимум 2000 слов. Пиши на русском языке. Формат — Markdown."
    )
    desc_part = f"\n\nДополнительные акценты от учителя: {description}" if description.strip() else ""
    user = (
        f"Напиши подробный конспект урока:\n\n"
        f"**Предмет:** {subject}\n**Класс:** {class_num} класс\n"
        f"**Тема:** {topic}\n**Учитель:** {teacher_name}\n**Школа:** {teacher_school}"
        f"{desc_part}\n\n"
        "## Структура (все разделы обязательны):\n\n"
        "### 1. Заголовок\nПредмет, класс, тема, ФИО учителя, дата.\n\n"
        "### 2. Цели урока\nОбразовательная, развивающая, воспитательная (по 2-3 предложения).\n\n"
        "### 3. Планируемые результаты по ФГОС\n"
        "Предметные, метапредметные (регулятивные, познавательные, коммуникативные), личностные.\n\n"
        "### 4. Оборудование и материалы\nПолный список с авторами учебников и ЭОР.\n\n"
        "### 5. Ход урока\n"
        "**5.1 Организационный момент (3 мин)** — слова учителя\n"
        "**5.2 Актуализация знаний (7 мин)** — 5-7 вопросов с ответами\n"
        "**5.3 Изучение нового материала (25 мин)** — ГЛАВНЫЙ раздел, подробно:\n"
        "все ключевые понятия с определениями, теория с объяснением, примеры, вопросы к классу\n"
        "**5.4 Первичное закрепление (8 мин)** — 3-4 задания с разборами\n"
        "**5.5 Итоги и рефлексия (5 мин)** — выводы и вопросы рефлексии\n"
        "**5.6 Домашнее задание** — с указанием параграфов, базовый и творческий уровень\n\n"
        "Пиши развёрнуто. Минимальный объём — 2000 слов."
    )
    return gigachat_chat(
        messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=8000,
    )


# ─── СБОРКА DOCX ─────────────────────────────────────────────────────────────

def _set_font(run, name="Times New Roman", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = DocxRGB(*color)
    # Кириллица
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def _set_para_spacing(para, before=0, after=6, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    if line:
        spacing.set(qn("w:line"), str(line))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def build_docx(md_text: str, subject: str, class_num: int, topic: str,
               teacher_name: str, teacher_school: str) -> bytes:
    doc = Document()

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # Титульный блок
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_para_spacing(title_para, after=2)
    run = title_para.add_run("КОНСПЕКТ УРОКА")
    _set_font(run, size=16, bold=True, color=(31, 73, 125))

    meta_lines = [
        f"Предмет: {subject}",
        f"Класс: {class_num}",
        f"Тема: {topic}",
        f"Учитель: {teacher_name}",
        f"Школа: {teacher_school}",
        f"Дата: _______________",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_para_spacing(p, after=0)
        run = p.add_run(line)
        _set_font(run, size=12)

    # Горизонтальная линия
    hr = doc.add_paragraph()
    _set_para_spacing(hr, before=60, after=60)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Парсим Markdown и добавляем содержимое
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Заголовки
        if line.startswith("#### "):
            p = doc.add_paragraph()
            _set_para_spacing(p, before=60, after=20)
            run = p.add_run(line[5:].strip())
            _set_font(run, size=11, bold=True, color=(79, 129, 189))
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_paragraph()
            _set_para_spacing(p, before=100, after=30)
            run = p.add_run(line[4:].strip())
            _set_font(run, size=12, bold=True, color=(31, 73, 125))
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_paragraph()
            _set_para_spacing(p, before=140, after=40)
            run = p.add_run(line[3:].strip())
            _set_font(run, size=14, bold=True, color=(17, 55, 100))
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            _set_para_spacing(p, before=140, after=60)
            run = p.add_run(line[2:].strip())
            _set_font(run, size=15, bold=True, color=(17, 55, 100))
            i += 1
            continue

        # Горизонтальный разделитель
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            _set_para_spacing(p, before=40, after=40)
            i += 1
            continue

        # Маркированный список
        if re.match(r"^[-*•]\s+", line):
            text = re.sub(r"^[-*•]\s+", "", line).strip()
            p = doc.add_paragraph(style="List Bullet")
            _set_para_spacing(p, before=0, after=0)
            _add_inline_markup(p, text, size=11)
            i += 1
            continue

        # Нумерованный список
        m = re.match(r"^(\d+)\.\s+(.+)", line)
        if m:
            text = m.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            _set_para_spacing(p, before=0, after=0)
            _add_inline_markup(p, text, size=11)
            i += 1
            continue

        # Пустая строка
        if not line.strip():
            i += 1
            continue

        # Обычный параграф
        p = doc.add_paragraph()
        _set_para_spacing(p, before=0, after=40, line=276)  # 1.15 spacing
        _add_inline_markup(p, line.strip(), size=12)
        i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline_markup(para, text: str, size: int = 12):
    """Добавляет текст с поддержкой **жирного** и *курсива* в параграф."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        # Текст до совпадения
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            _set_font(run, size=size)
        raw = m.group(0)
        if raw.startswith("**"):
            run = para.add_run(m.group(2))
            _set_font(run, size=size, bold=True)
        elif raw.startswith("*"):
            run = para.add_run(m.group(3))
            _set_font(run, size=size)
            run.font.italic = True
        elif raw.startswith("`"):
            run = para.add_run(m.group(4))
            _set_font(run, name="Courier New", size=size - 1)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        _set_font(run, size=size)


def safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "_", s)
    return s[:60].strip("_") or "konspekt"


# ─── HANDLER ─────────────────────────────────────────────────────────────────

def handler(event: dict, context) -> dict:
    """
    Генерирует конспект урока через GigaChat и возвращает DOCX файл.
    POST {subject, class_num, topic, description?, teacher_name, teacher_school, login?}
    -> {docx_b64, filename, word_count, topic, subject, class_num}
    """
    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": CORS, "body": ""}
    if event.get("httpMethod") != "POST":
        return _resp(405, {"error": "Method not allowed"})

    try:
        body = json.loads(event.get("body") or "{}")
    except Exception:
        return _resp(400, {"error": "Некорректный JSON"})

    subject = (body.get("subject") or "").strip()
    topic = (body.get("topic") or "").strip()
    teacher_name = (body.get("teacher_name") or "").strip()
    teacher_school = (body.get("teacher_school") or "").strip()
    description = (body.get("description") or "").strip()
    login = (body.get("login") or "").strip()

    try:
        class_num = int(body.get("class_num") or 0)
    except Exception:
        class_num = 0

    if not subject:
        return _resp(400, {"error": "subject обязателен"})
    if not topic:
        return _resp(400, {"error": "topic обязателен"})
    if class_num not in range(1, 12):
        return _resp(400, {"error": "class_num должен быть от 1 до 11"})

    # Списание токенов
    ok, err_msg = spend_ai_tokens(login, TOKENS_COST_SYNOPSIS)
    if not ok:
        return _resp(402, {"error": err_msg})

    md_text = gigachat_chat(
        messages=[
            {"role": "system", "content": (
                "Ты — опытный учитель-методист с 20-летним стажем, эксперт по ФГОС. "
                "Создавай профессиональные, ДЕТАЛЬНЫЕ конспекты уроков. "
                "Минимум 2000 слов. Пиши на русском языке. Формат — Markdown."
            )},
            {"role": "user", "content": (
                f"Напиши подробный конспект урока:\n\n"
                f"**Предмет:** {subject}\n**Класс:** {class_num} класс\n"
                f"**Тема:** {topic}\n**Учитель:** {teacher_name}\n**Школа:** {teacher_school}"
                + (f"\n\nДополнительные акценты: {description}" if description else "") +
                "\n\n## Структура (все разделы обязательны):\n\n"
                "### 1. Заголовок\nПредмет, класс, тема, ФИО учителя, дата.\n\n"
                "### 2. Цели урока\nОбразовательная, развивающая, воспитательная (по 2-3 предложения).\n\n"
                "### 3. Планируемые результаты по ФГОС\n"
                "Предметные, метапредметные, личностные (по 3-4 пункта).\n\n"
                "### 4. Оборудование и материалы\nПолный список с авторами учебников и ЭОР.\n\n"
                "### 5. Ход урока\n"
                "**5.1 Организационный момент (3 мин)**\n"
                "**5.2 Актуализация знаний (7 мин)** — 5-7 вопросов с ответами\n"
                "**5.3 Изучение нового материала (25 мин)** — ГЛАВНЫЙ раздел: все понятия с определениями, "
                "теория, примеры, вопросы к классу, что записать в тетрадь\n"
                "**5.4 Первичное закрепление (8 мин)** — 3-4 задания с разборами\n"
                "**5.5 Итоги и рефлексия (5 мин)**\n"
                "**5.6 Домашнее задание** — с параграфами, базовый и творческий уровень\n\n"
                "Минимум 2000 слов."
            )},
        ],
        max_tokens=8000,
    )

    word_count = len(md_text.split())

    docx_bytes = build_docx(
        md_text=md_text,
        subject=subject,
        class_num=class_num,
        topic=topic,
        teacher_name=teacher_name,
        teacher_school=teacher_school,
    )

    filename = f"Конспект_{safe_filename(topic)}_{class_num}кл.docx"

    return _resp(200, {
        "docx_b64": base64.b64encode(docx_bytes).decode(),
        "filename": filename,
        "word_count": word_count,
        "text": md_text,
        "topic": topic,
        "subject": subject,
        "class_num": class_num,
    })
