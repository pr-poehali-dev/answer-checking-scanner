"""
Генерация контрольной/проверочной работы или теста через GigaChat в .docx.
POST / body: {
  workType: "Тест" | "Проверочная работа" | "Контрольная работа",
  subject, classNum, topic, description?,
  part1Count: int (тестовые вопросы с вариантами),
  part2Count: int (открытые вопросы),
  teacherName, teacherSchool
}
Возвращает: {docx_b64, filename, workId, part1Count, part2Count, totalQuestions, answerKey, gradeScale, questions[]}
"""
import json
import os
import io
import re
import time
import base64
import urllib.request
import urllib.error

AUTH_URL = os.environ.get("AUTH_FUNCTION_URL", "https://functions.poehali.dev/b08ae7cf-6c0b-4178-acc9-4b62b2c2a61b")

TOKENS_COST_TEST = 3500


def spend_ai_tokens(login: str, amount: int) -> tuple[bool, str]:
    """Списывает токены через auth. Возвращает (ok, error_message)."""
    if not login:
        return True, ""
    try:
        req = urllib.request.Request(
            f"{AUTH_URL}?action=spend-tokens",
            data=json.dumps({"login": login, "amount": amount}).encode("utf-8"),
            method="POST",
            headers={"Content-Type": "application/json"},
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            data = json.loads(r.read().decode())
        return True, ""
    except urllib.error.HTTPError as e:
        err_body = {}
        try:
            err_body = json.loads(e.read().decode())
        except Exception:
            pass
        if e.code == 402:
            return False, err_body.get("error", "Недостаточно токенов")
        return True, ""
    except Exception:
        return True, ""


from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CORS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
}

def _resp(status: int, body: dict) -> dict:
    return {
        "statusCode": status,
        "headers": {**CORS, "Content-Type": "application/json"},
        "body": json.dumps(body, ensure_ascii=False),
        "isBase64Encoded": False,
    }


# ─── YANDEXGPT API ────────────────────────────────────────────────────────────

YANDEX_GPT_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"


def gigachat_chat(messages: list, max_tokens: int = 2400, temperature: float = 0.4,
                  req_timeout: int = 60, max_retries: int = 3) -> tuple[str, int]:
    api_key = os.environ.get("YANDEXGPT_API_KEY", "").strip()
    folder_id = os.environ.get("YANDEXGPT_FOLDER_ID", "").strip()
    if not api_key or not folder_id:
        raise RuntimeError("YANDEXGPT_API_KEY или YANDEXGPT_FOLDER_ID не заданы")
    yandex_messages = [{"role": m.get("role", "user"), "text": m.get("content", "")} for m in messages]
    payload = {
        "modelUri": f"gpt://{folder_id}/yandexgpt/latest",
        "completionOptions": {
            "stream": False,
            "temperature": temperature,
            "maxTokens": str(max_tokens),
        },
        "messages": yandex_messages,
    }
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            req = urllib.request.Request(
                YANDEX_GPT_URL,
                data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
                method="POST",
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Api-Key {api_key}",
                    "x-folder-id": folder_id,
                },
            )
            with urllib.request.urlopen(req, timeout=req_timeout) as r:
                body = json.loads(r.read().decode())
            alternatives = (body.get("result") or {}).get("alternatives") or []
            if not alternatives:
                raise RuntimeError(f"YandexGPT пустой ответ: {body}")
            text = alternatives[0].get("message", {}).get("text", "").strip()
            if not text:
                raise RuntimeError("YandexGPT вернул пустой текст")
            usage = (body.get("result") or {}).get("usage") or {}
            tokens_used = int(usage.get("totalTokens") or usage.get("completionTokens") or 0)
            return text, tokens_used
        except urllib.error.HTTPError as e:
            err_text = e.read().decode(errors="ignore")[:300]
            if e.code in (401, 403):
                raise RuntimeError(f"YandexGPT auth error {e.code}: {err_text}")
            last_err = RuntimeError(f"YandexGPT HTTP {e.code}: {err_text}")
            if attempt < max_retries:
                time.sleep(2.0)
        except Exception as e:
            last_err = RuntimeError(f"YandexGPT недоступен: {e}")
            if attempt < max_retries:
                time.sleep(2.0)
    raise last_err or RuntimeError("YandexGPT: не удалось получить ответ")


def extract_json(text: str) -> dict:
    text = text.strip()
    fence = re.search(r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", text)
    if fence:
        text = fence.group(1)
    else:
        s = text.find("{")
        e = text.rfind("}")
        if s >= 0 and e > s:
            text = text[s:e + 1]
    return json.loads(text)


# ─── ГЕНЕРАЦИЯ ВОПРОСОВ ─────────────────────────────────────────────────────

LETTERS = ["А", "Б", "В", "Г", "Д"]


def _generate_part1(work_type: str, subject: str, class_num: int, topic: str, description: str, count: int) -> tuple[list, int]:
    """Часть 1: вопросы с 4 вариантами ответа. Генерируется отдельным запросом."""
    if count <= 0:
        return [], 0
    system = (
        "Ты учитель-методист РФ. Создаёшь школьные тестовые вопросы строго по теме. "
        "Каждый вопрос имеет РОВНО 4 варианта ответа (А, Б, В, Г), из них 1 правильный. "
        "Возвращай ТОЛЬКО валидный JSON без markdown."
    )
    user = (
        f"Предмет: {subject}\n"
        f"Класс: {class_num}\n"
        f"Тема: {topic}\n"
        f"Описание: {description or '—'}\n"
        f"Тип работы: {work_type}\n\n"
        f"Составь РОВНО {count} тестовых вопросов с 4 вариантами ответа.\n"
        "Верни JSON строго в формате:\n"
        '{"questions": [\n'
        '  {"question": "Текст вопроса", "options": ["Вариант А", "Вариант Б", "Вариант В", "Вариант Г"], "answer": "А"},\n'
        '  ...\n'
        ']}\n'
        f"Требования:\n"
        f"- РОВНО {count} вопросов в массиве questions\n"
        f"- В каждом вопросе РОВНО 4 элемента в options\n"
        f"- answer — одна буква (А, Б, В или Г), указывающая правильный вариант\n"
        f"- Только 1 правильный ответ в каждом вопросе\n"
        f"- Вопросы разнообразные, проверяющие понимание темы"
    )
    max_tok = min(180 * count + 500, 3200)
    raw, _tok = gigachat_chat(
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=max_tok,
        temperature=0.5,
    )
    try:
        data = extract_json(raw)
    except Exception as e:
        raise RuntimeError(f"Не удалось разобрать вопросы части 1: {e}. Ответ: {raw[:300]}")

    items = data.get("questions") or data.get("part1") or []
    if not isinstance(items, list):
        items = []

    result = []
    for q in items[:count]:
        if not isinstance(q, dict):
            continue
        text = (q.get("question") or "").strip()
        opts = q.get("options") or []
        if not isinstance(opts, list):
            continue
        ans_raw = str(q.get("answer") or "").strip().upper()
        # Принимаем буквы А/Б/В/Г, латинские A/B/C/D, цифры 1-4
        lat_to_cyr = {"A": "А", "B": "Б", "C": "В", "D": "Г"}
        num_to_cyr = {"1": "А", "2": "Б", "3": "В", "4": "Г"}
        if ans_raw and ans_raw[0] in lat_to_cyr:
            ans = lat_to_cyr[ans_raw[0]]
        elif ans_raw and ans_raw[0] in num_to_cyr:
            ans = num_to_cyr[ans_raw[0]]
        elif ans_raw and ans_raw[0] in LETTERS:
            ans = ans_raw[0]
        else:
            ans = ""
        if not text or len(opts) < 2 or ans not in LETTERS:
            continue
        opts = [str(o).strip() for o in opts][:4]
        while len(opts) < 4:
            opts.append("—")
        if LETTERS.index(ans) >= len(opts):
            ans = "А"
        result.append({"question": text, "options": opts, "answer": ans})

    return result, _tok


def _generate_part2(work_type: str, subject: str, class_num: int, topic: str, description: str, count: int) -> tuple[list, int]:
    """Часть 2: открытые вопросы. Генерируется отдельным запросом."""
    if count <= 0:
        return [], 0
    system = (
        "Ты учитель-методист РФ. Создаёшь открытые вопросы школьной программы. "
        "Возвращай ТОЛЬКО валидный JSON без markdown."
    )
    user = (
        f"Предмет: {subject}\n"
        f"Класс: {class_num}\n"
        f"Тема: {topic}\n"
        f"Описание: {description or '—'}\n"
        f"Тип работы: {work_type}\n\n"
        f"Составь РОВНО {count} открытых вопросов (без вариантов ответа).\n"
        "Верни JSON строго в формате:\n"
        '{"questions": [\n'
        '  {"question": "Текст вопроса", "answer": "Краткий правильный ответ или решение"},\n'
        '  ...\n'
        ']}\n'
        f"РОВНО {count} элементов. Вопросы разнообразные."
    )
    max_tok = min(140 * count + 400, 3000)
    raw, _tok = gigachat_chat(
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=max_tok,
        temperature=0.5,
    )
    try:
        data = extract_json(raw)
    except Exception as e:
        raise RuntimeError(f"Не удалось разобрать вопросы части 2: {e}. Ответ: {raw[:300]}")

    items = data.get("questions") or data.get("part2") or []
    if not isinstance(items, list):
        items = []

    result = []
    for q in items[:count]:
        if not isinstance(q, dict):
            continue
        text = (q.get("question") or "").strip()
        ans = str(q.get("answer") or "").strip()
        if not text:
            continue
        result.append({"question": text, "answer": ans})

    return result, _tok


def generate_questions(work_type: str, subject: str, class_num: int, topic: str, description: str,
                       part1_count: int, part2_count: int) -> dict:
    """Запрашивает у GigaChat вопросы РАЗДЕЛЬНО для каждой части, чтобы гарантировать обе."""
    part1, tok1 = _generate_part1(work_type, subject, class_num, topic, description, part1_count)
    part2, tok2 = _generate_part2(work_type, subject, class_num, topic, description, part2_count)

    if part1_count > 0 and not part1:
        raise RuntimeError("ИИ не вернул валидных вопросов для части 1. Попробуйте ещё раз.")
    if part2_count > 0 and not part2:
        raise RuntimeError("ИИ не вернул валидных вопросов для части 2. Попробуйте ещё раз.")

    return {"part1": part1, "part2": part2, "total_tokens": tok1 + tok2}


# ─── DOCX BUILDER ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def build_docx(work_type: str, subject: str, class_num: int, topic: str,
               part1: list, part2: list, work_id: str,
               teacher_name: str, teacher_school: str,
               grade_scale: dict, max_score: int) -> bytes:
    doc = Document()

    # Стили базовые
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Поля страницы
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(1.5)

    # ── Шапка ──
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run(teacher_school or "Учебное заведение")
    run.bold = True
    run.font.size = Pt(11)

    # ── Заголовок работы ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(6)
    title.paragraph_format.space_after = Pt(2)
    r1 = title.add_run(f"{work_type.upper()}")
    r1.bold = True
    r1.font.size = Pt(16)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    r2 = sub.add_run(f"по предмету «{subject}» · {class_num} класс")
    r2.font.size = Pt(13)

    topic_p = doc.add_paragraph()
    topic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    topic_p.paragraph_format.space_after = Pt(10)
    r3 = topic_p.add_run(f"Тема: {topic}")
    r3.italic = True
    r3.font.size = Pt(12)

    # Метаданные таблицей: № работы, дата, ФИО ученика
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Light Grid Accent 1"
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_cells = meta.rows[0].cells
    meta_cells[0].text = "№ работы"
    meta_cells[1].text = "Дата"
    meta_cells[2].text = "Класс"
    meta_cells[3].text = "ФИО ученика"
    val_cells = meta.rows[1].cells
    val_cells[0].text = work_id
    val_cells[1].text = "____.____.________"
    val_cells[2].text = f"{class_num} _____"
    val_cells[3].text = "____________________________"
    for cell in meta_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        _set_cell_bg(cell, "E8EEF5")
    for cell in val_cells:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Инструкция ──
    instr = doc.add_paragraph()
    instr.paragraph_format.space_after = Pt(8)
    r = instr.add_run("Инструкция: ")
    r.bold = True
    parts_text = []
    if part1:
        parts_text.append(f"в части 1 ({len(part1)} зад.) выберите один правильный вариант ответа")
    if part2:
        parts_text.append(f"в части 2 ({len(part2)} зад.) дайте развёрнутый ответ")
    instr.add_run(". ".join(parts_text) + ". Время выполнения определяет учитель.")

    # ── Часть 1 ──
    if part1:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(4)
        rh = h.add_run(f"Часть 1. Тестовые задания с выбором ответа")
        rh.bold = True
        rh.font.size = Pt(13)

        for i, q in enumerate(part1, start=1):
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(4)
            qp.paragraph_format.space_after = Pt(2)
            qp.paragraph_format.left_indent = Cm(0)
            r = qp.add_run(f"{i}. ")
            r.bold = True
            qp.add_run(q["question"])

            for letter, opt in zip(LETTERS, q["options"]):
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Cm(0.8)
                op.paragraph_format.space_after = Pt(0)
                r1 = op.add_run(f"{letter})  ")
                r1.bold = True
                op.add_run(opt)

    # ── Часть 2 ──
    if part2:
        if part1:
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(4)
        rh = h.add_run(f"Часть 2. Задания с развёрнутым ответом")
        rh.bold = True
        rh.font.size = Pt(13)

        start = len(part1) + 1
        for i, q in enumerate(part2, start=0):
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(4)
            qp.paragraph_format.space_after = Pt(2)
            r = qp.add_run(f"{start + i}. ")
            r.bold = True
            qp.add_run(q["question"])
            # Линии для ответа
            for _ in range(3):
                lp = doc.add_paragraph("_" * 95)
                lp.paragraph_format.space_after = Pt(0)

    # ── Разрыв страницы перед ответами ──
    doc.add_page_break()

    # ── Ключи ответов ──
    key_h = doc.add_paragraph()
    key_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    key_h.paragraph_format.space_after = Pt(8)
    rk = key_h.add_run("ОТВЕТЫ (для учителя)")
    rk.bold = True
    rk.font.size = Pt(15)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.paragraph_format.space_after = Pt(10)
    info.add_run(f"{work_type} · {subject} · {class_num} класс · Тема: {topic} · № работы: {work_id}").italic = True

    if part1:
        h = doc.add_paragraph()
        h.paragraph_format.space_after = Pt(4)
        rh = h.add_run("Часть 1 (ответы)")
        rh.bold = True
        rh.font.size = Pt(13)

        # Таблица ответов части 1
        cols = min(10, len(part1))
        rows_needed = (len(part1) + cols - 1) // cols
        ans_table = doc.add_table(rows=rows_needed * 2, cols=cols)
        ans_table.style = "Light Grid Accent 1"
        for i in range(rows_needed):
            for j in range(cols):
                idx = i * cols + j
                if idx >= len(part1):
                    break
                num_cell = ans_table.cell(i * 2, j)
                ans_cell = ans_table.cell(i * 2 + 1, j)
                num_cell.text = str(idx + 1)
                ans_cell.text = part1[idx]["answer"]
                _set_cell_bg(num_cell, "E8EEF5")
                for p in num_cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                for p in ans_cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(11)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if part2:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(6)
        h.paragraph_format.space_after = Pt(4)
        rh = h.add_run("Часть 2 (примерные ответы)")
        rh.bold = True
        rh.font.size = Pt(13)

        start = len(part1) + 1
        for i, q in enumerate(part2):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{start + i}. ")
            r.bold = True
            p.add_run(q["answer"] or "—")

    # ── Шкала оценок ──
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    rh = h.add_run("Шкала оценивания")
    rh.bold = True
    rh.font.size = Pt(13)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    note.add_run(f"Максимум баллов: ").bold = False
    r = note.add_run(str(max_score))
    r.bold = True
    note.add_run(" (по 1 баллу за каждый правильный ответ)")

    scale_table = doc.add_table(rows=2, cols=5)
    scale_table.style = "Light Grid Accent 1"
    headers = scale_table.rows[0].cells
    headers[0].text = "«2»"
    headers[1].text = "«3»"
    headers[2].text = "«4»"
    headers[3].text = "«5»"
    headers[4].text = "Макс."
    vals = scale_table.rows[1].cells
    vals[0].text = f"0–{grade_scale['grade2'] - 1}" if grade_scale['grade2'] > 0 else "0"
    vals[1].text = f"{grade_scale['grade3']}–{grade_scale['grade4'] - 1}"
    vals[2].text = f"{grade_scale['grade4']}–{grade_scale['grade5'] - 1}"
    vals[3].text = f"{grade_scale['grade5']}–{max_score}"
    vals[4].text = str(max_score)
    for cell in headers:
        _set_cell_bg(cell, "E8EEF5")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
    for cell in vals:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(11)

    # ── Подпись учителя ──
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    foot.paragraph_format.space_before = Pt(20)
    if teacher_name:
        rf = foot.add_run(f"Составитель: {teacher_name}")
        rf.italic = True
        rf.font.size = Pt(11)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


# ─── ШКАЛА ОЦЕНОК ──────────────────────────────────────────────────────────

def compute_grade_scale(max_score: int) -> dict:
    """Стандартная школьная шкала: 50% — 3, 70% — 4, 90% — 5."""
    return {
        "grade1": 0,
        "grade2": max(1, round(max_score * 0.30)),
        "grade3": max(2, round(max_score * 0.50)),
        "grade4": max(3, round(max_score * 0.70)),
        "grade5": max(4, round(max_score * 0.90)),
    }


def safe_filename(s: str, max_len: int = 60) -> str:
    s = re.sub(r"[\\/:*?\"<>|]+", " ", s).strip()
    s = re.sub(r"\s+", " ", s)
    return s[:max_len] or "Работа"


def generate_work_id() -> str:
    """6-значный номер работы (как в appStore)."""
    import random
    return str(random.randint(100000, 999999))


# ─── HANDLER ───────────────────────────────────────────────────────────────

def handler(event: dict, context) -> dict:
    """Генерирует .docx-файл проверочной/контрольной работы или теста с помощью GigaChat."""
    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": CORS, "body": ""}

    method = event.get("httpMethod", "POST")
    if method != "POST":
        return _resp(405, {"error": "Метод не поддерживается"})

    body = {}
    raw = event.get("body") or ""
    if event.get("isBase64Encoded"):
        try:
            raw = base64.b64decode(raw).decode()
        except Exception:
            raw = ""
    try:
        body = json.loads(raw) if raw else {}
        if isinstance(body, str):
            body = json.loads(body)
    except Exception:
        body = {}

    work_type = (body.get("workType") or "Тест").strip()
    if work_type not in ("Тест", "Проверочная работа", "Контрольная работа"):
        work_type = "Тест"
    login = (body.get("login") or "").strip()
    subject = (body.get("subject") or "").strip()
    topic = (body.get("topic") or "").strip()
    description = (body.get("description") or "").strip()
    teacher_name = (body.get("teacherName") or "").strip()
    teacher_school = (body.get("teacherSchool") or "").strip()

    try:
        class_num = int(body.get("classNum") or 5)
    except (TypeError, ValueError):
        class_num = 5
    class_num = max(1, min(class_num, 11))

    try:
        part1_count = int(body.get("part1Count") or 0)
    except (TypeError, ValueError):
        part1_count = 0
    part1_count = max(0, min(part1_count, 30))

    try:
        part2_count = int(body.get("part2Count") or 0)
    except (TypeError, ValueError):
        part2_count = 0
    part2_count = max(0, min(part2_count, 10))

    if not subject:
        return _resp(400, {"error": "Укажите предмет"})
    if not topic:
        return _resp(400, {"error": "Укажите тему"})
    if part1_count + part2_count == 0:
        return _resp(400, {"error": "Должен быть хотя бы один вопрос (часть 1 или часть 2)"})

    # Проверяем лимит AI-запросов для trial-пользователей
    if login:
        try:
            limit_req = urllib.request.Request(
                f"{AUTH_URL}?action=check-ai-limit",
                data=json.dumps({"login": login}).encode("utf-8"),
                method="POST",
                headers={"Content-Type": "application/json"},
            )
            with urllib.request.urlopen(limit_req, timeout=10) as r:
                limit_data = json.loads(r.read().decode())
            if not limit_data.get("allowed"):
                return _resp(429, {"error": limit_data.get("error", "Достигнут лимит ИИ-запросов")})
        except urllib.error.HTTPError as e:
            err_body = json.loads(e.read().decode() or "{}")
            if e.code == 429:
                return _resp(429, {"error": err_body.get("error", "Достигнут лимит ИИ-запросов на сегодня")})
        except Exception:
            pass

    try:
        questions = generate_questions(work_type, subject, class_num, topic, description, part1_count, part2_count)
    except Exception as e:
        msg = str(e)
        if "timed out" in msg.lower() or "timeout" in msg.lower():
            return _resp(504, {"error": "Сервис GigaChat сейчас перегружен. Подождите минуту и попробуйте снова."})
        return _resp(500, {"error": f"Ошибка генерации вопросов: {msg}"})

    total_tokens_used = questions.get("total_tokens", 0)
    spend_ai_tokens(login, max(total_tokens_used, 1))

    part1 = questions["part1"]
    part2 = questions["part2"]
    actual_p1 = len(part1)
    actual_p2 = len(part2)
    total = actual_p1 + actual_p2
    max_score = total
    grade_scale = compute_grade_scale(max_score)

    work_id = generate_work_id()

    try:
        docx_bytes = build_docx(
            work_type=work_type,
            subject=subject,
            class_num=class_num,
            topic=topic,
            part1=part1,
            part2=part2,
            work_id=work_id,
            teacher_name=teacher_name,
            teacher_school=teacher_school,
            grade_scale=grade_scale,
            max_score=max_score,
        )
    except Exception as e:
        return _resp(500, {"error": f"Ошибка сборки .docx: {e}"})

    # Ключ ответов части 1 в формате "АБВГ..." (для appStore.work.answerKey)
    answer_key = "".join(q["answer"] for q in part1)

    filename = f"{safe_filename(work_type)} · {safe_filename(subject)} · {class_num} класс · {safe_filename(topic, 40)}.docx"

    return _resp(200, {
        "docx_b64": base64.b64encode(docx_bytes).decode(),
        "filename": filename,
        "size": len(docx_bytes),
        "workId": work_id,
        "workType": work_type,
        "subject": subject,
        "classNum": class_num,
        "topic": topic,
        "part1Count": actual_p1,
        "part2Count": actual_p2,
        "totalQuestions": total,
        "answerKey": answer_key,
        "maxScore": max_score,
        "gradeScale": grade_scale,
        "questions": {
            "part1": part1,
            "part2": part2,
        },
    })