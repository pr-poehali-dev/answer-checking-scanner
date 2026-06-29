"""
Генерация фирменного рабочего листа САОУ через YandexGPT в .docx.
Материалы подбираются по теме/классу/описанию строго по программе Минпросвещения РФ.
При необходимости ИИ добавляет иллюстрации (фото/карты) из Wikimedia.

POST / body: {
  subject, classNum, topic, description?,
  tasksCount: int (число заданий),
  withImages?: bool,
  teacherName, teacherSchool, login?
}
Возвращает: {docx_url?, docx_b64?, filename, size, title, tasks[], spent_rub, balance_rub}
"""
import json
import os
import io
import re
import time
import base64
import uuid
import urllib.request
import urllib.error
import urllib.parse

AUTH_URL = os.environ.get("AUTH_FUNCTION_URL", "https://functions.poehali.dev/b08ae7cf-6c0b-4178-acc9-4b62b2c2a61b")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CORS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
}

# Фирменные цвета САОУ
BRAND_DARK = "0D1B3E"
BRAND_BLUE = "00B4D8"
BRAND_LIGHT = "E8F4FB"
BRAND_ACCENT = "1B4F9C"


def _resp(status: int, body: dict) -> dict:
    return {
        "statusCode": status,
        "headers": {**CORS, "Content-Type": "application/json"},
        "body": json.dumps(body, ensure_ascii=False),
        "isBase64Encoded": False,
    }


def spend_ai_tokens(login: str, amount: int, action_label: str = "Рабочий лист") -> tuple[bool, str, float, float]:
    if not login:
        return True, "", 0.0, 0.0
    try:
        req = urllib.request.Request(
            f"{AUTH_URL}?action=spend-tokens",
            data=json.dumps({"login": login, "amount": amount, "action_label": action_label}).encode("utf-8"),
            method="POST",
            headers={"Content-Type": "application/json"},
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            resp = json.loads(r.read().decode())
        return True, "", float(resp.get("spent_rub") or 0), float(resp.get("balance_rub") or 0)
    except urllib.error.HTTPError as e:
        err_body = {}
        try:
            err_body = json.loads(e.read().decode())
        except Exception:
            pass
        if e.code == 402:
            return False, err_body.get("error", "Недостаточно средств"), 0.0, 0.0
        if e.code == 403:
            return False, err_body.get("error", "Для использования ИИ необходима активная подписка."), 0.0, 0.0
        return True, "", 0.0, 0.0
    except Exception:
        return True, "", 0.0, 0.0


# ─── YANDEXGPT ───────────────────────────────────────────────────────────────

YANDEX_GPT_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"


def yandex_chat(messages: list, max_tokens: int = 3000, temperature: float = 0.5,
                req_timeout: int = 90, max_retries: int = 3) -> tuple[str, int]:
    api_key = os.environ.get("YANDEXGPT_API_KEY", "").strip()
    folder_id = os.environ.get("YANDEXGPT_FOLDER_ID", "").strip()
    if not api_key or not folder_id:
        raise RuntimeError("YANDEXGPT_API_KEY или YANDEXGPT_FOLDER_ID не заданы")
    yandex_messages = [{"role": m.get("role", "user"), "text": m.get("content", "")} for m in messages]
    payload = {
        "modelUri": f"gpt://{folder_id}/yandexgpt/latest",
        "completionOptions": {"stream": False, "temperature": temperature, "maxTokens": str(max_tokens)},
        "messages": yandex_messages,
    }
    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            req = urllib.request.Request(
                YANDEX_GPT_URL,
                data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
                method="POST",
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Api-Key {api_key}",
                    "x-folder-id": folder_id,
                },
            )
            with urllib.request.urlopen(req, timeout=req_timeout) as r:
                body = json.loads(r.read().decode())
            alternatives = (body.get("result") or {}).get("alternatives") or []
            if not alternatives:
                raise RuntimeError(f"YandexGPT пустой ответ: {body}")
            text = alternatives[0].get("message", {}).get("text", "").strip()
            if not text:
                raise RuntimeError("YandexGPT вернул пустой текст")
            usage = (body.get("result") or {}).get("usage") or {}
            tokens_used = int(usage.get("totalTokens") or usage.get("completionTokens") or 0)
            return text, tokens_used
        except urllib.error.HTTPError as e:
            err_text = e.read().decode(errors="ignore")[:300]
            if e.code in (401, 403):
                raise RuntimeError(f"YandexGPT auth error {e.code}: {err_text}")
            last_err = RuntimeError(f"YandexGPT HTTP {e.code}: {err_text}")
            if attempt < max_retries:
                time.sleep(2.0)
        except Exception as e:
            last_err = RuntimeError(f"YandexGPT недоступен: {e}")
            if attempt < max_retries:
                time.sleep(2.0)
    raise last_err or RuntimeError("YandexGPT: не удалось получить ответ")


def _repair_truncated_json(text: str) -> str:
    stack, in_string, escape = [], False, False
    for ch in text:
        if escape:
            escape = False
            continue
        if ch == "\\":
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch in "{[":
            stack.append(ch)
        elif ch == "}":
            if stack and stack[-1] == "{":
                stack.pop()
        elif ch == "]":
            if stack and stack[-1] == "[":
                stack.pop()
    if in_string:
        text += '"'
    for ch in reversed(stack):
        text += "]" if ch == "[" else "}"
    return text


def extract_json(text: str) -> dict:
    text = text.strip()
    fence = re.search(r"```(?:json)?\s*(\{[\s\S]*)", text)
    if fence:
        text = fence.group(1)
        end = text.find("```")
        if end >= 0:
            text = text[:end]
    s = text.find("{")
    if s >= 0:
        text = text[s:]
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return json.loads(_repair_truncated_json(text))


# ─── ГЕНЕРАЦИЯ ЗАДАНИЙ ───────────────────────────────────────────────────────

def _clean_table(tbl) -> dict | None:
    """Нормализует таблицу-приложение: {headers: [...], rows: [[...], ...]}."""
    if not isinstance(tbl, dict):
        return None
    headers = tbl.get("headers") or []
    rows = tbl.get("rows") or []
    if not isinstance(headers, list) or not isinstance(rows, list):
        return None
    headers = [str(h).strip() for h in headers][:8]
    if not headers:
        return None
    clean_rows = []
    for row in rows[:12]:
        if not isinstance(row, list):
            continue
        cells = [str(c).strip() for c in row][:len(headers)]
        while len(cells) < len(headers):
            cells.append("")
        clean_rows.append(cells)
    if not clean_rows:
        return None
    return {"headers": headers, "rows": clean_rows}


def generate_worksheet_content(subject: str, class_num: int, topic: str,
                               description: str, tasks_count: int, with_images: bool) -> tuple[dict, int]:
    system = (
        "Ты опытный учитель-методист РФ. Составляешь учебные рабочие листы строго по "
        "Федеральной образовательной программе и материалам, утверждённым Министерством "
        "просвещения Российской Федерации (ФГОС). Содержание точное, без ошибок, "
        "соответствует возрасту и классу. Главный принцип: каждое задание должно быть "
        "ВЫПОЛНИМЫМ — если для решения нужны данные (числа, текст-источник, таблица, "
        "карта, изображение), ты обязан приложить эти данные прямо в задании, чтобы ученик "
        "мог выполнить его без посторонних источников. Возвращай ТОЛЬКО валидный JSON без markdown."
    )
    img_rule = (
        '- Если для выполнения задания НУЖНА иллюстрация (карта с данными, схема, фото объекта, '
        'репродукция, диаграмма) — добавь поле "image_query": короткий поисковый запрос на русском '
        'или английском (например "политическая карта России", "строение клетки растения схема", '
        '"портрет Пушкина"). Изображение должно быть напрямую связано с вопросом задания. '
        'Если иллюстрация не требуется — не добавляй это поле.\n'
        if with_images else
        '- Поле "image_query" не добавляй.\n'
    )
    user = (
        f"Предмет: {subject}\n"
        f"Класс: {class_num}\n"
        f"Тема: {topic}\n"
        f"Описание/акцент: {description or '—'}\n\n"
        f"Составь учебный рабочий лист по теме строго по программе Минпросвещения РФ для {class_num} класса.\n"
        f"Нужно РОВНО {tasks_count} заданий разного типа (вопросы, задачи, заполнить пропуски, "
        f"работа с таблицей данных, анализ текста-источника, работа с картой/изображением, "
        f"соотнести, проанализировать и т.п.).\n"
        "Верни JSON строго в формате:\n"
        '{\n'
        '  "title": "Краткое название рабочего листа",\n'
        '  "intro": "1-2 предложения вводной информации/мотивации по теме",\n'
        '  "tasks": [\n'
        '    {"number": 1, "type": "Тип задания", "instruction": "Текст задания для ученика", '
        '"content": "Текст-источник или данные для выполнения, если нужны словами; иначе пустая строка", '
        '"table": {"headers": ["Колонка 1", "Колонка 2"], "rows": [["знач", "знач"], ["знач", "знач"]]}, '
        '"answer_lines": 3, "image_query": "поисковый запрос для картинки (опционально)"}\n'
        '  ]\n'
        '}\n'
        "Требования:\n"
        f"- РОВНО {tasks_count} заданий в массиве tasks\n"
        "- instruction — понятная формулировка задания для ученика\n"
        "- ВАЖНО: задание должно быть выполнимо. Если для ответа нужны данные — приложи их.\n"
        "- Для заданий на анализ данных используй поле \"table\" с заголовками и строками "
        "(реальные осмысленные данные по теме). Если таблица не нужна — не добавляй поле \"table\".\n"
        "- В \"content\" клади текст-источник, набор фактов, условие задачи или ряд данных, "
        "когда ученику нужно на что-то опираться. Иначе оставь пустым.\n"
        "- Приложение (таблица, текст, изображение) и вопрос задания должны быть взаимосвязаны: "
        "вопрос задаётся именно по приложенным данным.\n"
        "- answer_lines — сколько пустых линий оставить для ответа (1-6, для устных 0)\n"
        "- Задания разнообразные и проверяют понимание темы\n"
        + img_rule +
        "- Только достоверная информация по программе РФ"
    )
    max_tok = min(300 * tasks_count + 1200, 8000)
    raw, tok = yandex_chat(
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        max_tokens=max_tok, temperature=0.45,
    )
    try:
        data = extract_json(raw)
    except Exception as e:
        raise RuntimeError(f"Не удалось разобрать задания: {e}. Ответ: {raw[:300]}")

    tasks = data.get("tasks") or []
    if not isinstance(tasks, list):
        tasks = []
    clean = []
    for i, t in enumerate(tasks[:tasks_count], start=1):
        if not isinstance(t, dict):
            continue
        instr = (t.get("instruction") or t.get("task") or "").strip()
        if not instr:
            continue
        try:
            lines = int(t.get("answer_lines", 3))
        except (TypeError, ValueError):
            lines = 3
        clean.append({
            "number": i,
            "type": (t.get("type") or "Задание").strip(),
            "instruction": instr,
            "content": (t.get("content") or "").strip(),
            "table": _clean_table(t.get("table")),
            "answer_lines": max(0, min(lines, 8)),
            "image_query": (t.get("image_query") or "").strip() if with_images else "",
        })
    if not clean:
        raise RuntimeError("ИИ не вернул валидных заданий. Попробуйте ещё раз.")
    return {
        "title": (data.get("title") or topic).strip(),
        "intro": (data.get("intro") or "").strip(),
        "tasks": clean,
    }, tok


# ─── ПОИСК ИЗОБРАЖЕНИЙ ───────────────────────────────────────────────────────

def fetch_wikimedia(query: str, timeout: int = 5) -> bytes | None:
    try:
        search_q = urllib.parse.quote(query)
        url = (
            f"https://commons.wikimedia.org/w/api.php"
            f"?action=query&generator=search&gsrnamespace=6"
            f"&gsrsearch={search_q}&gsrlimit=5"
            f"&prop=imageinfo&iiprop=url|mime|size"
            f"&iiurlwidth=700&format=json&origin=*"
        )
        req = urllib.request.Request(url, headers={"User-Agent": "SAOU-Edu-Bot/1.0"})
        with urllib.request.urlopen(req, timeout=timeout) as r:
            data = json.loads(r.read().decode())
        pages = (data.get("query") or {}).get("pages") or {}
        candidates = []
        for page in pages.values():
            ii = (page.get("imageinfo") or [{}])[0]
            mime = ii.get("mime", "")
            thumb_url = ii.get("thumburl") or ii.get("url", "")
            size = ii.get("size", 0)
            if mime in ("image/jpeg", "image/png") and thumb_url and 8000 < size < 5_000_000:
                candidates.append((size, thumb_url))
        candidates.sort(reverse=True)
        for _, thumb_url in candidates[:2]:
            try:
                req2 = urllib.request.Request(thumb_url, headers={"User-Agent": "SAOU-Edu-Bot/1.0"})
                with urllib.request.urlopen(req2, timeout=timeout) as r2:
                    return r2.read()
            except Exception:
                continue
        return None
    except Exception:
        return None


def fetch_images_for_tasks(tasks: list) -> dict:
    """Параллельно качает картинки для заданий с image_query. Возвращает {number: bytes}."""
    from concurrent.futures import ThreadPoolExecutor

    targets = [(t["number"], t["image_query"]) for t in tasks if t.get("image_query")]
    if not targets:
        return {}
    result = {}

    def work(item):
        num, q = item
        img = fetch_wikimedia(q)
        return num, img

    try:
        with ThreadPoolExecutor(max_workers=6) as ex:
            for num, img in ex.map(work, targets):
                if img:
                    result[num] = img
    except Exception:
        pass
    return result


# ─── S3 ──────────────────────────────────────────────────────────────────────

def upload_docx_to_s3(docx_bytes: bytes, filename: str) -> str:
    try:
        import boto3
        key_id = os.environ.get("AWS_ACCESS_KEY_ID", "")
        secret = os.environ.get("AWS_SECRET_ACCESS_KEY", "")
        if not key_id or not secret:
            return ""
        s3 = boto3.client(
            "s3",
            endpoint_url="https://bucket.poehali.dev",
            aws_access_key_id=key_id,
            aws_secret_access_key=secret,
        )
        safe = re.sub(r"[^a-zA-Z0-9._-]", "_", filename) or "worksheet.docx"
        key = f"worksheets/{uuid.uuid4().hex}_{safe}"
        s3.put_object(
            Bucket="files",
            Key=key,
            Body=docx_bytes,
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        return f"https://cdn.poehali.dev/projects/{key_id}/bucket/{key}"
    except Exception:
        return ""


# ─── DOCX BUILDER ────────────────────────────────────────────────────────────

def _set_cell_bg(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _page_border(section, color_hex: str):
    """Декоративная окантовка страницы."""
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "18")
        el.set(qn("w:space"), "24")
        el.set(qn("w:color"), color_hex)
        pg_borders.append(el)
    sect_pr.append(pg_borders)


def _no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def _add_data_table(doc, tbl: dict):
    """Рисует таблицу-приложение с данными: тёмная шапка + строки данных."""
    headers = tbl["headers"]
    rows = tbl["rows"]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        _set_cell_bg(hdr_cells[j], BRAND_ACCENT)
        p = hdr_cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        if i % 2 == 0:
            for c in cells:
                _set_cell_bg(c, BRAND_LIGHT)
        for j, val in enumerate(row):
            p = cells[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(10.5)


def build_docx(content: dict, subject: str, class_num: int, topic: str,
               teacher_name: str, teacher_school: str, images: dict) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.5)
    _page_border(section, BRAND_ACCENT)

    # ── Шапка: название учебного заведения (если указано) ──
    if teacher_school:
        head = doc.add_table(rows=1, cols=1)
        head.autofit = True
        hc = head.rows[0].cells[0]
        _set_cell_bg(hc, BRAND_DARK)
        hp = hc.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hp.add_run(teacher_school)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── Заголовок «РАБОЧИЙ ЛИСТ» ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(2)
    rt = title.add_run("РАБОЧИЙ ЛИСТ")
    rt.bold = True
    rt.font.size = Pt(20)
    rt.font.color.rgb = RGBColor.from_string(BRAND_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    rsub = subtitle.add_run(f"{subject} · {class_num} класс")
    rsub.font.size = Pt(13)
    rsub.bold = True
    rsub.font.color.rgb = RGBColor.from_string(BRAND_ACCENT)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(8)
    rn = name_p.add_run(content.get("title") or topic)
    rn.italic = True
    rn.font.size = Pt(12)

    # ── Поля ученика: ФИО и класс ──
    meta = doc.add_table(rows=1, cols=2)
    meta.style = "Table Grid"
    c0, c1 = meta.rows[0].cells
    c0.paragraphs[0].add_run("Ф.И.О. ученика: _______________________________").font.size = Pt(11)
    c1.paragraphs[0].add_run(f"Класс: {class_num} «____»     Дата: ___.___.20___").font.size = Pt(11)
    for c in (c0, c1):
        _set_cell_bg(c, BRAND_LIGHT)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True

    # ── Вводная ──
    intro = content.get("intro")
    if intro:
        ip = doc.add_paragraph()
        ip.paragraph_format.space_before = Pt(8)
        ip.paragraph_format.space_after = Pt(6)
        ri = ip.add_run(intro)
        ri.italic = True
        ri.font.size = Pt(11)
        ri.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

    # ── Задания ──
    for t in content["tasks"]:
        # Заголовок задания (плашка-номер + тип)
        head_p = doc.add_paragraph()
        head_p.paragraph_format.space_before = Pt(10)
        head_p.paragraph_format.space_after = Pt(2)
        rnum = head_p.add_run(f"Задание {t['number']}")
        rnum.bold = True
        rnum.font.size = Pt(13)
        rnum.font.color.rgb = RGBColor.from_string(BRAND_ACCENT)
        if t.get("type"):
            rtype = head_p.add_run(f"   ·   {t['type']}")
            rtype.font.size = Pt(10)
            rtype.italic = True
            rtype.font.color.rgb = RGBColor(0x77, 0x88, 0x99)

        # Формулировка
        instr_p = doc.add_paragraph()
        instr_p.paragraph_format.space_after = Pt(2)
        instr_p.add_run(t["instruction"]).font.size = Pt(12)

        # Приложение: текст-источник / данные
        if t.get("content"):
            src_label = doc.add_paragraph()
            src_label.paragraph_format.space_before = Pt(2)
            src_label.paragraph_format.space_after = Pt(0)
            rl = src_label.add_run("Приложение к заданию:")
            rl.bold = True
            rl.font.size = Pt(10)
            rl.font.color.rgb = RGBColor.from_string(BRAND_ACCENT)

            cont_p = doc.add_paragraph()
            cont_p.paragraph_format.left_indent = Cm(0.5)
            cont_p.paragraph_format.space_after = Pt(3)
            rc = cont_p.add_run(t["content"])
            rc.font.size = Pt(11)
            rc.italic = True

        # Приложение: таблица данных
        tbl = t.get("table")
        if tbl and tbl.get("headers"):
            _add_data_table(doc, tbl)

        # Иллюстрация
        img_bytes = images.get(t["number"])
        if img_bytes:
            try:
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(2)
                img_p.paragraph_format.space_after = Pt(2)
                img_p.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(8.5))
            except Exception:
                pass

        # Линии для ответа
        for _ in range(t.get("answer_lines", 0)):
            lp = doc.add_paragraph("_" * 92)
            _no_space(lp)
            lp.paragraph_format.space_after = Pt(2)

    # ── Подвал: подпись учителя (если указана) ──
    if teacher_name:
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.paragraph_format.space_before = Pt(14)
        rf = footer.add_run(f"Учитель: {teacher_name}")
        rf.font.size = Pt(9)
        rf.italic = True
        rf.font.color.rgb = RGBColor(0x66, 0x77, 0x88)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def safe_filename(s: str, max_len: int = 60) -> str:
    s = re.sub(r"[\\/:*?\"<>|]+", " ", s).strip()
    s = re.sub(r"\s+", " ", s)
    return s[:max_len] or "Рабочий лист"


# ─── HANDLER ─────────────────────────────────────────────────────────────────

def handler(event: dict, context) -> dict:
    """Генерирует фирменный рабочий лист САОУ (.docx) с помощью ИИ по программе Минпросвещения РФ."""
    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": CORS, "body": ""}

    method = event.get("httpMethod", "POST")
    if method != "POST":
        return _resp(405, {"error": "Метод не поддерживается"})

    raw = event.get("body") or ""
    if event.get("isBase64Encoded"):
        try:
            raw = base64.b64decode(raw).decode()
        except Exception:
            raw = ""
    try:
        body = json.loads(raw) if raw else {}
        if isinstance(body, str):
            body = json.loads(body)
    except Exception:
        body = {}

    login = (body.get("login") or "").strip()
    subject = (body.get("subject") or "").strip()
    topic = (body.get("topic") or "").strip()
    description = (body.get("description") or "").strip()
    teacher_name = (body.get("teacherName") or "").strip()
    teacher_school = (body.get("teacherSchool") or "").strip()
    with_images = bool(body.get("withImages", True))

    try:
        class_num = int(body.get("classNum") or 5)
    except (TypeError, ValueError):
        class_num = 5
    class_num = max(1, min(class_num, 11))

    try:
        tasks_count = int(body.get("tasksCount") or 5)
    except (TypeError, ValueError):
        tasks_count = 5
    tasks_count = max(1, min(tasks_count, 20))

    if not subject:
        return _resp(400, {"error": "Укажите предмет"})
    if not topic:
        return _resp(400, {"error": "Укажите тему"})

    # Лимит AI-запросов
    if login:
        try:
            limit_req = urllib.request.Request(
                f"{AUTH_URL}?action=check-ai-limit",
                data=json.dumps({"login": login}).encode("utf-8"),
                method="POST",
                headers={"Content-Type": "application/json"},
            )
            with urllib.request.urlopen(limit_req, timeout=10) as r:
                limit_data = json.loads(r.read().decode())
            if not limit_data.get("allowed"):
                return _resp(429, {"error": limit_data.get("error", "Достигнут лимит ИИ-запросов")})
        except urllib.error.HTTPError as e:
            err_body = json.loads(e.read().decode() or "{}")
            if e.code == 429:
                return _resp(429, {"error": err_body.get("error", "Достигнут лимит ИИ-запросов на сегодня")})
        except Exception:
            pass

    try:
        content, tokens_used = generate_worksheet_content(
            subject, class_num, topic, description, tasks_count, with_images
        )
    except Exception as e:
        msg = str(e)
        if "timed out" in msg.lower() or "timeout" in msg.lower():
            return _resp(504, {"error": "ИИ-сервис сейчас перегружен. Подождите минуту и попробуйте снова."})
        return _resp(500, {"error": f"Ошибка генерации заданий: {msg}"})

    _, _, spent_rub, balance_rub = spend_ai_tokens(login, max(tokens_used, 1))

    images = {}
    if with_images:
        try:
            images = fetch_images_for_tasks(content["tasks"])
        except Exception:
            images = {}

    try:
        docx_bytes = build_docx(content, subject, class_num, topic, teacher_name, teacher_school, images)
    except Exception as e:
        return _resp(500, {"error": f"Ошибка сборки .docx: {e}"})

    filename = f"Рабочий лист · {safe_filename(subject)} · {class_num} класс · {safe_filename(topic, 40)}.docx"
    docx_url = upload_docx_to_s3(docx_bytes, filename)

    resp_body = {
        "docx_url": docx_url,
        "filename": filename,
        "size": len(docx_bytes),
        "title": content.get("title") or topic,
        "subject": subject,
        "classNum": class_num,
        "topic": topic,
        "tasksCount": len(content["tasks"]),
        "withImages": with_images,
        "imagesAdded": len(images),
        "tasks": content["tasks"],
        "intro": content.get("intro", ""),
        "spent_rub": spent_rub,
        "balance_rub": balance_rub,
    }
    if not docx_url or len(docx_bytes) < 3_000_000:
        resp_body["docx_b64"] = base64.b64encode(docx_bytes).decode()
    return _resp(200, resp_body)