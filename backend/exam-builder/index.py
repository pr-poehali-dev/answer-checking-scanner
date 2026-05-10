"""
Экзамены ФИПИ: генерация полного варианта ОГЭ или ЕГЭ из встроенного банка заданий.
Без ИИ. Мгновенно. Каждое задание выбирается случайно из готовой коллекции
заданий по соответствующему номеру в структуре ФИПИ.

POST / body:
{
  "examType": "ОГЭ" | "ЕГЭ",
  "subject": str,
  "teacherName": str,
  "teacherSchool": str,
  "variantNum": int (optional)
}
Возвращает: { docx_b64, answers_docx_b64, filename, answers_filename, ... }

GET ?action=subjects&examType=ОГЭ|ЕГЭ -> список предметов
"""
import json
import io
import re
import base64
import random

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from bank import EXAM_BANK

CORS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
}


def _resp(status: int, body: dict) -> dict:
    return {
        "statusCode": status,
        "headers": {**CORS, "Content-Type": "application/json"},
        "body": json.dumps(body, ensure_ascii=False),
        "isBase64Encoded": False,
    }


# ─── DOCX HELPERS ────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rPr.insert(0, rFonts)


def add_para(doc, text, size=12, bold=False, italic=False, align=None, color=None, indent_cm=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


# ─── ВЫБОР ЗАДАНИЙ ──────────────────────────────────────────────────────────

def build_variant_tasks(exam_type: str, subject: str, variant_num: int = 0) -> list:
    """Берёт банк по предмету, для каждого слота выбирает случайное задание.
    variant_num задаёт seed, чтобы один и тот же номер варианта давал
    одинаковый набор заданий, а разные номера — разные."""
    info = EXAM_BANK.get(exam_type, {}).get(subject)
    if not info:
        return []
    rng = random.Random(f"{exam_type}|{subject}|{variant_num}")
    tasks = []
    for slot in info["structure"]:
        num = slot["num"]
        bank = info["tasks"].get(num, [])
        if not bank:
            picked = {
                "question": f"[Задание по теме «{slot['topic']}». Решите согласно инструкции.]",
                "options": [],
                "answer": "—",
                "explanation": f"Тема: {slot['topic']}",
            }
        else:
            picked = rng.choice(bank)
        tasks.append({
            "num": num,
            "type": slot["type"],
            "topic": slot["topic"],
            "points": slot["points"],
            "instruction": slot["instruction"],
            "question": picked.get("question", ""),
            "options": picked.get("options", []),
            "answer": str(picked.get("answer", "")),
            "explanation": picked.get("explanation", ""),
        })
    return tasks


# ─── DOCX: ВАРИАНТ (СТРОГИЙ ФОРМАТ ФИПИ) ────────────────────────────────────

EXAM_DURATION = {
    ("ОГЭ", "Русский язык"): "3 часа 55 минут (235 минут)",
    ("ОГЭ", "Математика"): "3 часа 55 минут (235 минут)",
    ("ОГЭ", "Физика"): "3 часа (180 минут)",
    ("ОГЭ", "Химия"): "3 часа (180 минут)",
    ("ОГЭ", "Биология"): "3 часа (180 минут)",
    ("ОГЭ", "История"): "3 часа (180 минут)",
    ("ОГЭ", "Обществознание"): "3 часа (180 минут)",
    ("ОГЭ", "География"): "2 часа 30 минут (150 минут)",
    ("ОГЭ", "Информатика"): "2 часа 30 минут (150 минут)",
    ("ОГЭ", "Иностранный язык"): "2 часа 15 минут (135 минут)",
    ("ЕГЭ", "Русский язык"): "3 часа 30 минут (210 минут)",
    ("ЕГЭ", "Математика (профиль)"): "3 часа 55 минут (235 минут)",
    ("ЕГЭ", "Математика (база)"): "3 часа (180 минут)",
    ("ЕГЭ", "Физика"): "3 часа 55 минут (235 минут)",
    ("ЕГЭ", "Химия"): "3 часа 30 минут (210 минут)",
    ("ЕГЭ", "Биология"): "3 часа 55 минут (235 минут)",
    ("ЕГЭ", "История"): "3 часа 30 минут (210 минут)",
    ("ЕГЭ", "Обществознание"): "3 часа 30 минут (210 минут)",
    ("ЕГЭ", "География"): "3 часа (180 минут)",
    ("ЕГЭ", "Информатика"): "3 часа 55 минут (235 минут)",
    ("ЕГЭ", "Иностранный язык"): "3 часа 10 минут (190 минут)",
    ("ЕГЭ", "Литература"): "3 часа 55 минут (235 минут)",
}


def _add_horizontal_line(doc):
    """Тонкая горизонтальная линия как в бланках ФИПИ."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _instructions_text(exam_type: str, subject: str, total_tasks: int, total_points: int) -> str:
    duration = EXAM_DURATION.get((exam_type, subject), "регламент ФИПИ")
    return (
        f"Экзаменационная работа состоит из {total_tasks} заданий. "
        f"На выполнение работы по предмету «{subject}» отводится {duration}. "
        "Ответы к заданиям записываются в виде числа, слова, последовательности букв или цифр в отведённое поле. "
        "Если в задании предусмотрен развёрнутый ответ, запишите полное решение и обоснование. "
        "При выполнении работы не разрешается пользоваться учебниками, рабочими тетрадями, "
        "справочниками, мобильными телефонами и другими средствами связи. "
        "При необходимости можно пользоваться черновиком. Записи в черновике "
        "не учитываются при оценивании работы. Баллы, полученные за выполненные задания, суммируются. "
        f"Максимальный первичный балл за работу — {total_points}. "
        "Постарайтесь выполнить как можно больше заданий и набрать наибольшее количество баллов. "
        "После завершения работы проверьте, чтобы ответ к каждому заданию был записан "
        "под правильным номером. Желаем успеха!"
    )


def build_variant_docx(exam_type: str, subject: str, tasks: list,
                       teacher_name: str, teacher_school: str,
                       variant_num: int) -> bytes:
    doc = Document()

    # Стандартные поля официальных бланков ФИПИ
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    total_tasks = len(tasks)
    total_points = sum(t["points"] for t in tasks)

    # ── ТИТУЛЬНЫЙ ЛИСТ ────────────────────────────────────────────────────
    if teacher_school:
        add_para(doc, teacher_school, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_horizontal_line(doc)
    doc.add_paragraph()

    add_para(doc, "Единый государственный экзамен" if exam_type == "ЕГЭ"
             else "Основной государственный экзамен",
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "по предмету", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"«{subject.upper()}»", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, f"ВАРИАНТ № {variant_num}", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # Метка задания (имитация шапки бланка ФИПИ)
    add_para(doc, "Бланк ответов № 1", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph()
    _add_horizontal_line(doc)

    # Поля учащегося
    add_para(doc, "Фамилия: ____________________________________________________", size=11)
    add_para(doc, "Имя: ________________________________________________________", size=11)
    add_para(doc, "Отчество: ___________________________________________________", size=11)
    add_para(doc, "Класс: __________   Дата: «____» ______________ 20___ г.", size=11)
    if teacher_name:
        add_para(doc, f"Учитель: {teacher_name}", size=11, italic=True)
    _add_horizontal_line(doc)
    doc.add_paragraph()

    # ── ИНСТРУКЦИЯ ────────────────────────────────────────────────────────
    add_para(doc, "Инструкция по выполнению работы", size=12, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, _instructions_text(exam_type, subject, total_tasks, total_points),
             size=11)
    doc.add_paragraph()

    # Разрыв страницы перед заданиями
    from docx.enum.text import WD_BREAK
    p_break = doc.add_paragraph()
    p_break.add_run().add_break(WD_BREAK.PAGE)
    doc.add_paragraph()
    add_para(doc, "ЭКЗАМЕНАЦИОННАЯ РАБОТА", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"{exam_type} · {subject} · Вариант № {variant_num}",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x55, 0x55, 0x55))
    _add_horizontal_line(doc)
    doc.add_paragraph()

    # ── ЗАДАНИЯ ───────────────────────────────────────────────────────────
    current_part = None
    part_titles = {
        "A": ("ЧАСТЬ 1",
              "Ответами к заданиям являются число, последовательность цифр или слово. "
              "Запишите ответ в отведённое поле."),
        "B": ("ЧАСТЬ 2",
              "Ответом к заданиям является краткий ответ в виде числа, слова или "
              "последовательности символов."),
        "C": ("ЧАСТЬ 3",
              "Для записи ответов используйте отведённое поле. Запишите сначала номер "
              "задания, а затем полное обоснованное решение и ответ."),
    }

    for task in tasks:
        t_type = task["type"]
        target_part = "A" if t_type in ("choice", "multi") else \
                      "B" if t_type == "short" else "C"

        if target_part != current_part:
            current_part = target_part
            doc.add_paragraph()
            title, hint = part_titles[target_part]
            add_para(doc, title, size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     color=RGBColor(0x00, 0x00, 0x00))
            add_para(doc, hint, size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     color=RGBColor(0x55, 0x55, 0x55))
            _add_horizontal_line(doc)
            doc.add_paragraph()

        # Заголовок задания: «1.» жирно слева, тема справа курсивом
        p_num = doc.add_paragraph()
        run_num = p_num.add_run(f"{task['num']}.  ")
        set_font(run_num, size=13, bold=True)
        run_topic = p_num.add_run(task["topic"])
        set_font(run_topic, size=10, italic=True, color=RGBColor(0x55, 0x55, 0x55))

        # Инструкция к заданию
        add_para(doc, task["instruction"], size=11)

        if task.get("question"):
            add_para(doc, task["question"], size=11)

        # Варианты ответов
        if task.get("options"):
            for opt in task["options"]:
                add_para(doc, str(opt), size=11, indent_cm=0.7)

        # Поле для ответа
        if t_type in ("choice", "multi", "short"):
            ans_p = doc.add_paragraph()
            run_ans = ans_p.add_run("Ответ: ")
            set_font(run_ans, size=11, bold=True)
            run_field = ans_p.add_run("│" + " " * 30 + "│")
            set_font(run_field, size=11, color=RGBColor(0x00, 0x00, 0x00))
        elif t_type in ("long", "essay"):
            add_para(doc, "Запишите номер задания и полное решение:", size=10,
                     italic=True, color=RGBColor(0x55, 0x55, 0x55))
            for _ in range(8):
                line_p = doc.add_paragraph()
                pPr = line_p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "888888")
                pBdr.append(bottom)
                pPr.append(pBdr)
                set_font(line_p.add_run(" "), size=11)

        doc.add_paragraph()

    # ── ФИНАЛЬНАЯ СТРАНИЦА ────────────────────────────────────────────────
    doc.add_paragraph()
    _add_horizontal_line(doc)
    add_para(doc, f"Конец работы. Максимальный балл: {total_points}.",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x55, 0x55, 0x55))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── DOCX: ОТВЕТЫ ───────────────────────────────────────────────────────────

def build_answers_docx(exam_type: str, subject: str, tasks: list,
                       teacher_name: str, variant_num: int) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    add_para(doc, f"ОТВЕТЫ К ВАРИАНТУ {variant_num}",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"{exam_type}  ·  {subject}",
             size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, f"Составил: {teacher_name}",
             size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    add_para(doc, "ОТВЕТЫ И КРИТЕРИИ ОЦЕНИВАНИЯ", size=12, bold=True)
    doc.add_paragraph()

    for task in tasks:
        add_para(doc, f"Задание {task['num']}. {task['topic']}  [{task['points']} б.]",
                 size=12, bold=True)
        if task.get("answer"):
            add_para(doc, f"Ответ: {task['answer']}", size=12, color=RGBColor(0x0A, 0x66, 0x0A))
        if task.get("explanation"):
            add_para(doc, f"Пояснение: {task['explanation']}",
                     size=11, italic=True, color=RGBColor(0x44, 0x44, 0x44))
        doc.add_paragraph()

    total_points = sum(t["points"] for t in tasks)
    doc.add_paragraph()
    add_para(doc, "ШКАЛА ПЕРЕВОДА БАЛЛОВ", size=12, bold=True)
    add_para(doc, f"Максимальный балл: {total_points}", size=12)
    add_para(doc, f"«5» — {int(total_points * 0.85)}–{total_points} баллов", size=12)
    add_para(doc, f"«4» — {int(total_points * 0.70)}–{int(total_points * 0.84)} баллов", size=12)
    add_para(doc, f"«3» — {int(total_points * 0.50)}–{int(total_points * 0.69)} баллов", size=12)
    add_para(doc, f"«2» — 0–{int(total_points * 0.49)} баллов", size=12)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── HANDLER ────────────────────────────────────────────────────────────────

def handler(event: dict, context) -> dict:
    """
    Генерация полного варианта ОГЭ/ЕГЭ из банка заданий ФИПИ — без ИИ.

    GET ?action=subjects&examType=ОГЭ|ЕГЭ -> список предметов.
    POST { examType, subject, teacherName, teacherSchool, variantNum? } -> docx_b64, answers_docx_b64.
    """
    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200,
                "headers": {**CORS, "Content-Type": "application/json"},
                "body": ""}

    if event.get("httpMethod") == "GET":
        qs = event.get("queryStringParameters") or {}
        action = qs.get("action", "")
        if action == "subjects":
            exam_type = qs.get("examType", "ОГЭ")
            subjects = sorted(EXAM_BANK.get(exam_type, {}).keys())
            return _resp(200, {"subjects": subjects})
        return _resp(200, {
            "oge_subjects": sorted(EXAM_BANK.get("ОГЭ", {}).keys()),
            "ege_subjects": sorted(EXAM_BANK.get("ЕГЭ", {}).keys()),
        })

    if event.get("httpMethod") != "POST":
        return _resp(405, {"error": "Метод не поддерживается"})

    try:
        body = json.loads(event.get("body") or "{}")
    except Exception:
        return _resp(400, {"error": "Некорректный JSON"})

    exam_type = (body.get("examType") or "").strip()
    subject = (body.get("subject") or "").strip()
    teacher_name = (body.get("teacherName") or "Учитель").strip()
    teacher_school = (body.get("teacherSchool") or "").strip()
    variant_num = body.get("variantNum")
    try:
        variant_num = int(variant_num) if variant_num else random.randint(1, 99)
    except Exception:
        variant_num = random.randint(1, 99)

    if exam_type not in ("ОГЭ", "ЕГЭ"):
        return _resp(400, {"error": "Укажите тип экзамена: ОГЭ или ЕГЭ"})
    if not subject:
        return _resp(400, {"error": "Укажите предмет"})

    info = EXAM_BANK.get(exam_type, {}).get(subject)
    if not info:
        available = sorted(EXAM_BANK.get(exam_type, {}).keys())
        return _resp(400, {
            "error": f"Предмет «{subject}» недоступен для {exam_type}. Доступны: {', '.join(available)}"
        })

    tasks = build_variant_tasks(exam_type, subject, variant_num)
    if not tasks:
        return _resp(500, {"error": "В банке нет заданий по этому предмету"})

    variant_bytes = build_variant_docx(exam_type, subject, tasks,
                                       teacher_name, teacher_school, variant_num)
    answers_bytes = build_answers_docx(exam_type, subject, tasks,
                                       teacher_name, variant_num)

    safe_subject = re.sub(r"[^\w\-]", "_", subject)
    filename = f"{exam_type}_{safe_subject}_вариант_{variant_num:02d}.docx"
    answers_filename = f"{exam_type}_{safe_subject}_ответы_{variant_num:02d}.docx"
    total_points = sum(t["points"] for t in tasks)

    return _resp(200, {
        "docx_b64": base64.b64encode(variant_bytes).decode(),
        "answers_docx_b64": base64.b64encode(answers_bytes).decode(),
        "filename": filename,
        "answers_filename": answers_filename,
        "examType": exam_type,
        "subject": subject,
        "variantNum": variant_num,
        "totalTasks": len(tasks),
        "totalPoints": total_points,
        "size": len(variant_bytes),
    })