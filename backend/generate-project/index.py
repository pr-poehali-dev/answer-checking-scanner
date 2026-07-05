"""
Генерация индивидуального проекта / реферата / курсовой / доклада / сочинения / текста
через ИИ (YandexGPT) с оформлением по стандартам РФ (Минобрнауки / Минпросвещения).
Возвращает DOCX и PDF (base64).

POST / body: {
  work_type: "project"|"referat"|"coursework"|"report"|"essay"|"text",
  topic, subject?, description?, author_name?, school?, login?
}
-> {docx_b64, pdf_b64, filename, text, word_count, page_estimate, spent_rub, balance_rub}
"""
import json
import os
import io
import re
import base64
import urllib.request
import urllib.error
from datetime import datetime

import psycopg2
import boto3
from docx import Document
from docx.shared import Pt, Cm, RGBColor as DocxRGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AUTH_URL = os.environ.get("AUTH_FUNCTION_URL", "https://functions.poehali.dev/b08ae7cf-6c0b-4178-acc9-4b62b2c2a61b")
SCHEMA = os.environ.get("MAIN_DB_SCHEMA", "t_p31556921_answer_checking_scan")


def get_conn():
    return psycopg2.connect(os.environ["DATABASE_URL"])


def s3_client():
    return boto3.client(
        "s3",
        endpoint_url="https://bucket.poehali.dev",
        aws_access_key_id=os.environ["AWS_ACCESS_KEY_ID"],
        aws_secret_access_key=os.environ["AWS_SECRET_ACCESS_KEY"],
    )


def upload_to_s3(data: bytes, key: str, content_type: str) -> str:
    s3 = s3_client()
    s3.put_object(Bucket="files", Key=key, Body=data, ContentType=content_type)
    return f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{key}"

CORS = {
    "Access-Control-Allow-Origin": "*",
    "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
    "Access-Control-Allow-Headers": "Content-Type",
}

# ─── Типы работ: требования по объёму (Минобр РФ) ────────────────────────────
WORK_TYPES = {
    "project":    {"label": "Индивидуальный проект", "min_pages": 10, "max_pages": 20, "min_words": 2800, "sections": True},
    "referat":    {"label": "Реферат",               "min_pages": 10, "max_pages": 15, "min_words": 2800, "sections": True},
    "coursework": {"label": "Курсовая работа",        "min_pages": 20, "max_pages": 25, "min_words": 5600, "sections": True},
    "report":     {"label": "Доклад",                "min_pages": 10, "max_pages": 15, "min_words": 2800, "sections": True},
    "essay":      {"label": "Сочинение",              "min_pages": 1,  "max_pages": 2,  "min_words": 300,  "sections": False},
    "text":       {"label": "Текст",                 "min_pages": 15, "max_pages": 20, "min_words": 4200, "sections": True},
}

# ~280 слов на страницу (Times New Roman 14, полуторный интервал, поля по ГОСТ)
WORDS_PER_PAGE = 280


def _resp(status: int, body: dict) -> dict:
    return {
        "statusCode": status,
        "headers": {**CORS, "Content-Type": "application/json"},
        "body": json.dumps(body, ensure_ascii=False),
        "isBase64Encoded": False,
    }


# ─── ТОКЕНЫ ───────────────────────────────────────────────────────────────────

def spend_ai_tokens(login: str, amount: int, action_label: str) -> tuple[bool, str, float, float]:
    if not login:
        return True, "", 0.0, 0.0
    try:
        req = urllib.request.Request(
            f"{AUTH_URL}?action=spend-tokens",
            data=json.dumps({"login": login, "amount": amount, "action_label": action_label}).encode("utf-8"),
            method="POST",
            headers={"Content-Type": "application/json"},
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            resp = json.loads(r.read().decode())
        return True, "", float(resp.get("spent_rub") or 0), float(resp.get("balance_rub") or 0)
    except urllib.error.HTTPError as e:
        err_body = {}
        try:
            err_body = json.loads(e.read().decode())
        except Exception:
            pass
        if e.code == 402:
            return False, err_body.get("error", "Недостаточно средств"), 0.0, 0.0
        if e.code == 403:
            return False, err_body.get("error", "Для использования ИИ необходима активная подписка."), 0.0, 0.0
        return True, "", 0.0, 0.0
    except Exception:
        return True, "", 0.0, 0.0


# ─── ИИ API (YandexGPT) ───────────────────────────────────────────────────────

YANDEX_GPT_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/completion"


def ai_chat(messages: list, max_tokens: int = 8000, temperature: float = 0.6,
            req_timeout: int = 160) -> tuple[str, int]:
    api_key = os.environ.get("YANDEXGPT_API_KEY", "").strip()
    folder_id = os.environ.get("YANDEXGPT_FOLDER_ID", "").strip()
    if not api_key or not folder_id:
        raise RuntimeError("YANDEXGPT_API_KEY или YANDEXGPT_FOLDER_ID не заданы")

    yandex_messages = [{"role": m.get("role", "user"), "text": m.get("content", "")} for m in messages]
    payload = {
        "modelUri": f"gpt://{folder_id}/yandexgpt/latest",
        "completionOptions": {"stream": False, "temperature": temperature, "maxTokens": str(max_tokens)},
        "messages": yandex_messages,
    }
    req = urllib.request.Request(
        YANDEX_GPT_URL,
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        method="POST",
        headers={"Content-Type": "application/json",
                 "Authorization": f"Api-Key {api_key}", "x-folder-id": folder_id},
    )
    with urllib.request.urlopen(req, timeout=req_timeout) as r:
        body = json.loads(r.read().decode())
    alternatives = (body.get("result") or {}).get("alternatives") or []
    if not alternatives:
        raise RuntimeError(f"ИИ вернул пустой ответ: {body}")
    text = alternatives[0].get("message", {}).get("text", "").strip()
    usage = (body.get("result") or {}).get("usage") or {}
    tokens_used = int(usage.get("totalTokens") or usage.get("completionTokens") or 0)
    return text, tokens_used


# ─── ГЕНЕРАЦИЯ СТРУКТУРЫ (план глав) ─────────────────────────────────────────

SYSTEM_BASE = (
    "Ты — эксперт по академическому письму, полностью соблюдающий требования "
    "Министерства науки и высшего образования РФ и Министерства просвещения РФ, "
    "стандарты ГОСТ 7.32 и ФГОС. Ты пишешь оригинальные, уникальные, авторские "
    "тексты научно-учебного стиля на русском языке. Текст должен легко проходить "
    "проверку на антиплагиат: используй разнообразные формулировки, избегай клише "
    "и прямого копирования источников, перефразируй факты своими словами. "
    "НИКОГДА не упоминай, что текст создан учеником, студентом или искусственным "
    "интеллектом. Пиши от нейтрального авторского лица. Формат вывода — Markdown."
)


def generate_outline(work: dict, topic: str, subject: str, description: str) -> list:
    """Генерирует список глав (заголовков) для работы."""
    if not work["sections"]:
        return []
    desc = f"\nПожелания автора: {description}" if description.strip() else ""
    sub = f"\nПредмет/дисциплина: {subject}" if subject.strip() else ""
    user = (
        f"Составь план (оглавление) для работы типа «{work['label']}» по теме: «{topic}».{sub}{desc}\n\n"
        "Верни СТРОГО JSON без пояснений в формате:\n"
        '{"chapters": ["Введение", "Глава 1. ...", "1.1 ...", "Глава 2. ...", "Заключение", "Список литературы"]}\n\n'
        f"Работа должна содержать введение, {2 if work['min_pages'] < 20 else 3}-4 основные главы "
        "с подпунктами, заключение и список литературы. Заголовки — по теме, конкретные и осмысленные."
    )
    raw, _ = ai_chat(
        messages=[{"role": "system", "content": SYSTEM_BASE},
                  {"role": "user", "content": user}],
        max_tokens=1500, temperature=0.5,
    )
    chapters = _parse_json_chapters(raw)
    if not chapters:
        # запасной универсальный план
        chapters = ["Введение", f"Глава 1. Теоретические основы темы «{topic}»",
                    f"Глава 2. Практический анализ темы «{topic}»", "Заключение", "Список литературы"]
    return chapters


def _parse_json_chapters(raw: str) -> list:
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    if not m:
        return []
    try:
        data = json.loads(m.group(0))
        ch = data.get("chapters") or []
        return [str(c).strip() for c in ch if str(c).strip()]
    except Exception:
        return []


def generate_chapter(work: dict, topic: str, subject: str, chapter: str,
                     min_words: int, context_titles: list) -> str:
    """Генерирует содержимое одной главы/раздела."""
    is_intro = "введени" in chapter.lower()
    is_concl = "заключени" in chapter.lower()
    is_refs = "литератур" in chapter.lower() or "источник" in chapter.lower()

    if is_refs:
        user = (
            f"Составь список литературы (не менее 8 источников) по теме «{topic}» "
            f"для работы «{work['label']}». Оформи по ГОСТ Р 7.0.100-2018: авторы, название, "
            "город, издательство, год, страницы. Используй реальные типы источников "
            "(учебники, научные статьи, монографии, электронные ресурсы). Верни нумерованный список."
        )
        text, _ = ai_chat(messages=[{"role": "system", "content": SYSTEM_BASE},
                                     {"role": "user", "content": user}], max_tokens=2000, temperature=0.4)
        return text

    plan_note = ""
    if context_titles:
        plan_note = "\nОбщий план работы: " + "; ".join(context_titles) + "."

    if is_intro:
        focus = ("Напиши развёрнутое введение: актуальность темы, проблема, объект и предмет, "
                 "цель и задачи, методы, практическая значимость.")
    elif is_concl:
        focus = ("Напиши заключение: основные выводы по каждой задаче, итоговый вывод, "
                 "практическая значимость и перспективы. Без нумерации задач в лоб.")
    else:
        focus = ("Раскрой раздел подробно и содержательно: определения ключевых понятий, "
                 "теория, факты, примеры, анализ, при необходимости — сравнения и аргументы. "
                 "Пиши связным научным текстом абзацами, без воды.")

    sub = f" по предмету «{subject}»" if subject.strip() else ""
    user = (
        f"Пишем работу «{work['label']}» на тему «{topic}»{sub}.{plan_note}\n\n"
        f"Напиши раздел: «{chapter}».\n{focus}\n\n"
        f"Объём этого раздела — не менее {min_words} слов. "
        "Пиши оригинальным авторским текстом, уникально, чтобы проходило антиплагиат. "
        "Не используй заголовок повторно, не пиши мета-комментарии. Только содержимое раздела в Markdown."
    )
    text, _ = ai_chat(messages=[{"role": "system", "content": SYSTEM_BASE},
                                {"role": "user", "content": user}],
                      max_tokens=8000, temperature=0.6)
    return text


def generate_simple(work: dict, topic: str, subject: str, description: str) -> str:
    """Для сочинения/короткого текста — единый запрос."""
    desc = f"\nПожелания автора: {description}" if description.strip() else ""
    sub = f"\nПредмет: {subject}" if subject.strip() else ""
    if work["label"] == "Сочинение":
        task = (
            f"Напиши сочинение на тему «{topic}».{sub}{desc}\n\n"
            f"Объём — не менее {work['min_words']} слов. Структура: вступление, "
            "основная часть с аргументами и примерами, вывод. Оригинальный авторский текст, "
            "грамотный литературный русский язык, уникальный для антиплагиата."
        )
    else:
        task = (
            f"Напиши развёрнутый текст на тему «{topic}».{sub}{desc}\n\n"
            f"Объём — не менее {work['min_words']} слов. Логичная структура, "
            "оригинальный авторский стиль, уникальный для антиплагиата."
        )
    text, _ = ai_chat(messages=[{"role": "system", "content": SYSTEM_BASE},
                                {"role": "user", "content": task}], max_tokens=8000, temperature=0.7)
    return text


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def _set_font(run, name="Times New Roman", size=14, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = DocxRGB(*color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.insert(0, rFonts)


def _spacing(para, before=0, after=6, line=360):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    if line:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _inline(para, text: str, size=14):
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            _set_font(para.add_run(text[pos:m.start()]), size=size)
        raw = m.group(0)
        if raw.startswith("**"):
            _set_font(para.add_run(m.group(2)), size=size, bold=True)
        elif raw.startswith("*"):
            _set_font(para.add_run(m.group(3)), size=size, italic=True)
        elif raw.startswith("`"):
            _set_font(para.add_run(m.group(4)), name="Courier New", size=size - 1)
        pos = m.end()
    if pos < len(text):
        _set_font(para.add_run(text[pos:]), size=size)


def build_docx(work: dict, topic: str, subject: str, author_name: str, school: str,
               chapters: list, bodies: list, simple_text: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)

    # Титульный лист
    def centered(text, size=14, bold=False, before=0, after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p, before=before, after=after)
        _set_font(p.add_run(text), size=size, bold=bold)
        return p

    centered(school or "Образовательное учреждение", size=12, before=0, after=200)
    for _ in range(5):
        centered("", size=12)
    centered(work["label"].upper(), size=16, bold=True, after=120)
    centered(f"на тему: «{topic}»", size=14, after=40)
    if subject:
        centered(f"по дисциплине: {subject}", size=12, after=40)
    for _ in range(6):
        centered("", size=12)
    if author_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _spacing(p, after=20)
        _set_font(p.add_run(f"Выполнил(а): {author_name}"), size=12)
    for _ in range(4):
        centered("", size=12)
    import datetime
    centered(str(datetime.datetime.now().year), size=12)

    doc.add_page_break()

    # Содержание (если есть главы)
    if chapters:
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(h, after=120)
        _set_font(h.add_run("СОДЕРЖАНИЕ"), size=14, bold=True)
        for ch in chapters:
            p = doc.add_paragraph()
            _spacing(p, after=40)
            _set_font(p.add_run(ch), size=14)
        doc.add_page_break()

    # Тело
    if simple_text:
        _render_markdown(doc, simple_text)
    else:
        for ch, body in zip(chapters, bodies):
            hp = doc.add_paragraph()
            _spacing(hp, before=120, after=80)
            _set_font(hp.add_run(ch), size=14, bold=True)
            _render_markdown(doc, body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_markdown(doc, md_text: str):
    for raw_line in md_text.split("\n"):
        line = raw_line.rstrip()
        if line.startswith("### ") or line.startswith("## ") or line.startswith("# ") or line.startswith("#### "):
            title = line.lstrip("#").strip()
            p = doc.add_paragraph()
            _spacing(p, before=100, after=40)
            _set_font(p.add_run(title), size=14, bold=True, color=(31, 55, 100))
            continue
        if line.strip() in ("---", "***", "___"):
            continue
        if re.match(r"^[-*•]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _spacing(p, after=20)
            _inline(p, re.sub(r"^[-*•]\s+", "", line).strip())
            continue
        m = re.match(r"^(\d+)[.)]\s+(.+)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _spacing(p, after=20)
            _inline(p, m.group(2).strip())
            continue
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        _spacing(p, after=40, line=360)
        _inline(p, line.strip())


# ─── PDF (reportlab) ─────────────────────────────────────────────────────────

def build_pdf(work: dict, topic: str, subject: str, author_name: str, school: str,
              chapters: list, bodies: list, simple_text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_RIGHT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Регистрируем шрифт с кириллицей
    font_name = "DejaVu"
    try:
        for path in ("/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
                     "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"):
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont(font_name, path))
                break
        else:
            font_name = "Helvetica"
    except Exception:
        font_name = "Helvetica"

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm,
                            leftMargin=3 * cm, rightMargin=1.5 * cm)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("body", parent=styles["Normal"], fontName=font_name,
                                fontSize=14, leading=21, alignment=TA_JUSTIFY, firstLineIndent=1.25 * cm, spaceAfter=6)
    h_style = ParagraphStyle("h", parent=styles["Normal"], fontName=font_name, fontSize=14,
                             leading=20, spaceBefore=10, spaceAfter=6, textColor="#1f3764")
    center = ParagraphStyle("c", parent=styles["Normal"], fontName=font_name, fontSize=14,
                            leading=20, alignment=TA_CENTER)
    right = ParagraphStyle("r", parent=styles["Normal"], fontName=font_name, fontSize=12,
                           leading=18, alignment=TA_RIGHT)

    def esc(t):
        return (t or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    story = []
    story.append(Paragraph(esc(school or "Образовательное учреждение"), center))
    story.append(Spacer(1, 6 * cm))
    story.append(Paragraph(f"<b>{esc(work['label'].upper())}</b>", ParagraphStyle("t", parent=center, fontSize=18)))
    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(f"на тему: «{esc(topic)}»", center))
    if subject:
        story.append(Paragraph(f"по дисциплине: {esc(subject)}", center))
    story.append(Spacer(1, 4 * cm))
    if author_name:
        story.append(Paragraph(f"Выполнил(а): {esc(author_name)}", right))
    import datetime
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph(str(datetime.datetime.now().year), center))
    story.append(PageBreak())

    if chapters:
        story.append(Paragraph("<b>СОДЕРЖАНИЕ</b>", center))
        story.append(Spacer(1, 0.5 * cm))
        for ch in chapters:
            story.append(Paragraph(esc(ch), body_style))
        story.append(PageBreak())

    def render_md(md):
        for line in md.split("\n"):
            line = line.rstrip()
            if not line.strip():
                continue
            if line.lstrip().startswith("#"):
                story.append(Paragraph("<b>" + esc(line.lstrip("#").strip()) + "</b>", h_style))
                continue
            if line.strip() in ("---", "***", "___"):
                continue
            clean = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", line)
            clean = re.sub(r"\*(.+?)\*", r"<i>\1</i>", clean)
            clean = re.sub(r"^[-*•]\s+", "• ", clean)
            story.append(Paragraph(esc_keep_tags(clean), body_style))

    def esc_keep_tags(t):
        # экранируем &, но сохраняем <b>/<i>
        t = t.replace("&", "&amp;")
        return t

    if simple_text:
        render_md(simple_text)
    else:
        for ch, body in zip(chapters, bodies):
            story.append(Paragraph("<b>" + esc(ch) + "</b>", h_style))
            render_md(body)

    doc.build(story)
    return buf.getvalue()


def safe_filename(s: str) -> str:
    s = re.sub(r'[\\/:*?"<>|]', "_", s)
    return s[:60].strip("_") or "rabota"


# ─── HANDLER ─────────────────────────────────────────────────────────────────

def handler(event: dict, context) -> dict:
    """Генерирует индивидуальную работу (проект/реферат/курсовая/доклад/сочинение/текст) в DOCX и PDF."""
    if event.get("httpMethod") == "OPTIONS":
        return {"statusCode": 200, "headers": CORS, "body": ""}

    qs = event.get("queryStringParameters") or {}
    action = (qs.get("action") or "").strip().lower()

    # ── GET my-works — история работ ученика ──────────────────────────────────
    if event.get("httpMethod") == "GET" and action == "my-works":
        login = (qs.get("login") or "").strip()
        if not login:
            return _resp(400, {"error": "Укажите login"})
        conn = get_conn()
        try:
            cur = conn.cursor()
            cur.execute(
                f"""SELECT id, work_type, work_label, topic, subject, word_count,
                           page_estimate, docx_url, pdf_url, created_at
                    FROM {SCHEMA}.project_works WHERE author_login = %s
                    ORDER BY created_at DESC LIMIT 100""",
                (login,),
            )
            items = [{
                "id": x[0], "work_type": x[1], "work_label": x[2], "topic": x[3],
                "subject": x[4], "word_count": x[5], "page_estimate": x[6],
                "docx_url": x[7], "pdf_url": x[8], "created_at": str(x[9]),
            } for x in cur.fetchall()]
            return _resp(200, {"items": items})
        finally:
            conn.close()

    if event.get("httpMethod") != "POST":
        return _resp(405, {"error": "Method not allowed"})

    try:
        body = json.loads(event.get("body") or "{}")
    except Exception:
        return _resp(400, {"error": "Некорректный JSON"})

    work_type = (body.get("work_type") or "").strip()
    topic = (body.get("topic") or "").strip()
    subject = (body.get("subject") or "").strip()
    description = (body.get("description") or "").strip()
    work = WORK_TYPES.get(work_type)
    if not work:
        return _resp(400, {"error": "Неизвестный тип работы"})
    if not topic:
        return _resp(400, {"error": "Укажите тему работы"})

    # ── ШАГ 1: план глав ─────────────────────────────────────────────────────
    if action == "outline":
        if work["sections"]:
            chapters = generate_outline(work, topic, subject, description)
        else:
            chapters = []  # сочинение/текст — единым куском
        return _resp(200, {"chapters": chapters, "work_label": work["label"], "sections": work["sections"]})

    # ── ШАГ 2: одна глава (или весь простой текст) ───────────────────────────
    if action == "chapter":
        if not work["sections"]:
            text = generate_simple(work, topic, subject, description)
            return _resp(200, {"chapter": "", "body": text})
        chapter = (body.get("chapter") or "").strip()
        all_chapters = body.get("all_chapters") or []
        if not chapter:
            return _resp(400, {"error": "chapter обязателен"})
        content_chapters = [c for c in all_chapters if "литератур" not in c.lower() and "источник" not in c.lower()]
        per = max(350, int(work["min_words"] / max(1, len(content_chapters))))
        text = generate_chapter(work, topic, subject, chapter, per, all_chapters)
        return _resp(200, {"chapter": chapter, "body": text})

    # ── ШАГ 3: сборка файлов + сохранение ────────────────────────────────────
    if action == "build":
        author_name = (body.get("author_name") or "").strip()
        school = (body.get("school") or "").strip()
        login = (body.get("login") or "").strip()
        chapters = body.get("chapters") or []
        bodies = body.get("bodies") or []
        simple_text = (body.get("simple_text") or "").strip()

        full_text = simple_text or "\n\n".join(f"{c}\n{b}" for c, b in zip(chapters, bodies))
        word_count = len(full_text.split())
        page_estimate = max(work["min_pages"], round(word_count / WORDS_PER_PAGE))

        tokens = max(3000, int(word_count * 2.2))
        _, _, spent_rub, balance_rub = spend_ai_tokens(login, tokens, work["label"])

        docx_bytes = build_docx(work, topic, subject, author_name, school, chapters, bodies, simple_text)
        try:
            pdf_bytes = build_pdf(work, topic, subject, author_name, school, chapters, bodies, simple_text)
            pdf_b64 = base64.b64encode(pdf_bytes).decode()
        except Exception as e:
            print(f"[generate-project] PDF error: {e}")
            pdf_b64 = None

        filename = f"{work['label']}_{safe_filename(topic)}"

        docx_url = None
        pdf_url = None
        try:
            ts = int(datetime.utcnow().timestamp())
            base_key = f"projects/{login or 'anon'}/{ts}_{safe_filename(topic)}"
            docx_url = upload_to_s3(
                docx_bytes, f"{base_key}.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            if pdf_b64:
                pdf_url = upload_to_s3(base64.b64decode(pdf_b64), f"{base_key}.pdf", "application/pdf")
            if login:
                conn = get_conn()
                try:
                    cur = conn.cursor()
                    cur.execute(
                        f"""INSERT INTO {SCHEMA}.project_works
                            (author_login, work_type, work_label, topic, subject,
                             word_count, page_estimate, docx_url, pdf_url)
                            VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)""",
                        (login, work_type, work["label"], topic, subject or None,
                         word_count, page_estimate, docx_url, pdf_url),
                    )
                    conn.commit()
                finally:
                    conn.close()
        except Exception as e:
            print(f"[generate-project] save history error: {e}")

        return _resp(200, {
            "docx_b64": base64.b64encode(docx_bytes).decode(),
            "pdf_b64": pdf_b64,
            "docx_url": docx_url,
            "pdf_url": pdf_url,
            "filename": filename,
            "text": full_text,
            "chapters": chapters,
            "word_count": word_count,
            "page_estimate": page_estimate,
            "work_label": work["label"],
            "topic": topic,
            "spent_rub": spent_rub,
            "balance_rub": balance_rub,
        })

    return _resp(400, {"error": "Неизвестное действие. Используйте action=outline|chapter|build"})